from pathlib import Path
from docx import Document
from docx2pdf import convert


BASE_DIR = Path(__file__).resolve().parent

TEMPLATE_FILE = BASE_DIR / "template" / "Scorecard_Template.docx"
OUTPUT_DIR = BASE_DIR / "output"

OUTPUT_DIR.mkdir(exist_ok=True)


def replace_text_in_paragraph(paragraph, replacements):
    """
    Replace placeholders even when Word splits them
    across multiple runs.
    """

    for old_text, new_text in replacements.items():

        while old_text in paragraph.text:

            runs = paragraph.runs

            full_text = "".join(
                run.text for run in runs
            )

            start = full_text.find(old_text)

            if start == -1:
                break

            end = start + len(old_text)

            current_position = 0

            start_run_index = None
            end_run_index = None

            start_offset = None
            end_offset = None

            # Find starting and ending runs
            for i, run in enumerate(runs):

                run_start = current_position
                run_end = current_position + len(run.text)

                if (
                    start_run_index is None
                    and run_start <= start < run_end
                ):

                    start_run_index = i
                    start_offset = start - run_start

                if (
                    run_start < end <= run_end
                ):

                    end_run_index = i
                    end_offset = end - run_start
                    break

                current_position = run_end

            if (
                start_run_index is None
                or end_run_index is None
            ):
                break

            # Same run
            if start_run_index == end_run_index:

                run = runs[start_run_index]

                run.text = (
                    run.text[:start_offset]
                    + str(new_text)
                    + run.text[end_offset:]
                )

            else:

                start_run = runs[start_run_index]
                end_run = runs[end_run_index]

                # Text before placeholder
                before = start_run.text[:start_offset]

                # Text after placeholder
                after = end_run.text[end_offset:]

                # Put replacement into first run
                start_run.text = (
                    before
                    + str(new_text)
                )

                # Remove placeholder from middle/end runs
                for i in range(
                    start_run_index + 1,
                    end_run_index
                ):
                    runs[i].text = ""

                # Keep text after placeholder
                end_run.text = after


def replace_text_in_table(table, replacements):
    """
    Replace placeholders inside Word tables.
    """

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(
                    paragraph,
                    replacements
                )


def create_scorecard(
    year,
    academic_year,
    name,
    application_number,
    rank,
    score
):
    """
    Create one personalized scorecard.
    """

    # Load template
    document = Document(TEMPLATE_FILE)

    # Dynamic values
    replacements = {
        "{{ACADEMIC_YEAR}}": academic_year,
        "{{NAME}}": name,
        "{{APPLICATION_NUMBER}}": application_number,
        "{{RANK}}": rank,
        "{{SCORE}}": score,
    }

    # Replace in paragraphs
    for paragraph in document.paragraphs:

        replace_text_in_paragraph(
            paragraph,
            replacements
        )

    # Replace in tables
    for table in document.tables:

        replace_text_in_table(
            table,
            replacements
        )

    # Temporary DOCX
    safe_name = name.strip().replace(" ", "_")

    docx_file = OUTPUT_DIR / f"{safe_name}_Scorecard.docx"
    pdf_file = OUTPUT_DIR / f"{safe_name}_Scorecard.pdf"

    # Save personalized Word document
    document.save(docx_file)

    # Convert DOCX → PDF
    convert(
        str(docx_file),
        str(pdf_file)
    )

    print("\n" + "=" * 70)
    print("SCORECARD GENERATED")
    print("=" * 70)

    print(f"Name: {name}")
    print(f"Application Number: {application_number}")
    print(f"Rank: {rank}")
    print(f"Score: {score}/100")

    print(f"\nPDF: {pdf_file}")

    return pdf_file


# --------------------------------------------------
# TEST WITH SHIVAM GULJANI
# --------------------------------------------------

if __name__ == "__main__":
    create_scorecard(
        year="2023",
        academic_year="2023-24",
        name="Shivam Guljani",
        application_number="VGU_2023_55_035055",
        rank="VGUCET775",
        score="40"
    )