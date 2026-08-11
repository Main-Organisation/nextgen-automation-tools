from pathlib import Path
import pandas as pd
from docx import Document
from docx2pdf import convert


# ============================================================
# PATH CONFIGURATION
# ============================================================

BASE_DIR = Path(__file__).resolve().parent

EXCEL_FILE = (
    BASE_DIR
    / "generated"
    / "MCA_Students_Scores_and_Ranks_Updated.xlsx"
)

TEMPLATE_FILE = (
    BASE_DIR
    / "template"
    / "Scorecard_Template.docx"
)

OUTPUT_DIR = BASE_DIR / "output"

TEMP_DOCX_DIR = BASE_DIR / "generated" / "temp_docx"

LOG_FILE = (
    BASE_DIR
    / "generated"
    / "scorecard_generation_log.xlsx"
)


# Create required folders
OUTPUT_DIR.mkdir(exist_ok=True)
TEMP_DOCX_DIR.mkdir(exist_ok=True)


# ============================================================
# TEXT REPLACEMENT
# ============================================================

def replace_text_in_paragraph(paragraph, replacements):
    """
    Replace placeholders even when Word splits
    the placeholder across multiple runs.
    """

    full_text = "".join(
        run.text for run in paragraph.runs
    )

    if not full_text:
        return

    new_text = full_text

    for old_text, replacement in replacements.items():

        new_text = new_text.replace(
            old_text,
            str(replacement)
        )

    # Nothing changed
    if new_text == full_text:
        return

    # Put the complete modified text into
    # the first run and clear remaining runs.
    if paragraph.runs:

        paragraph.runs[0].text = new_text

        for run in paragraph.runs[1:]:

            run.text = ""

def replace_text_in_table(table, replacements):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                replace_text_in_paragraph(
                    paragraph,
                    replacements
                )


# ============================================================
# CREATE ONE SCORECARD
# ============================================================

def create_scorecard(
    year,
    academic_year,
    name,
    application_number,
    rank,
    score
):

    document = Document(TEMPLATE_FILE)

    replacements = {
        "{{ACADEMIC_YEAR}}": academic_year,
        "{{NAME}}": name,
        "{{APPLICATION_NUMBER}}": application_number,
        "{{RANK}}": rank,
        "{{SCORE}}": score,
    }


    # Replace placeholders in paragraphs
    for paragraph in document.paragraphs:

        replace_text_in_paragraph(
            paragraph,
            replacements
        )


    # Replace placeholders in tables
    for table in document.tables:

        replace_text_in_table(
            table,
            replacements
        )


    # --------------------------------------------------------
    # File name
    # --------------------------------------------------------

    safe_name = (
        name
        .strip()
        .replace("/", "_")
        .replace("\\", "_")
        .replace(" ", "_")
    )


    pdf_file = (
        OUTPUT_DIR
        / f"{safe_name}_Scorecard.pdf"
    )


    # If same name already exists,
    # use Application Number to avoid overwriting
    if pdf_file.exists():

        pdf_file = (
            OUTPUT_DIR
            / f"{safe_name}_{application_number}_Scorecard.pdf"
        )


    # Temporary Word file
    docx_file = (
        TEMP_DOCX_DIR
        / f"{safe_name}_{application_number}.docx"
    )


    # Save temporary DOCX
    document.save(docx_file)


    # Convert Word → PDF
    convert(
        str(docx_file),
        str(pdf_file)
    )


    # Remove temporary DOCX
    if docx_file.exists():

        docx_file.unlink()


    return pdf_file


# ============================================================
# MAIN BULK PROCESS
# ============================================================

def generate_all_scorecards():

    print("\n" + "=" * 70)
    print("VGU SCORECARD BULK GENERATOR")
    print("=" * 70)


    # --------------------------------------------------------
    # Check files
    # --------------------------------------------------------

    if not EXCEL_FILE.exists():

        print(
            f"❌ Excel file not found:\n{EXCEL_FILE}"
        )

        return


    if not TEMPLATE_FILE.exists():

        print(
            f"❌ Template file not found:\n{TEMPLATE_FILE}"
        )

        return


    # --------------------------------------------------------
    # Read Excel
    # --------------------------------------------------------

    print("\nReading Excel...")

    df = pd.read_excel(
        EXCEL_FILE
    )


    print(
        f"Records found: {len(df)}"
    )


    # --------------------------------------------------------
    # Validate required columns
    # --------------------------------------------------------

    required_columns = [
        "Year",
        "Name",
        "Score",
        "Rank",
        "Application Number"
    ]


    missing_columns = [
        column
        for column in required_columns
        if column not in df.columns
    ]


    if missing_columns:

        print(
            "\n❌ Missing columns:"
        )

        for column in missing_columns:

            print(
                f"   - {column}"
            )

        return


    # --------------------------------------------------------
    # Generation log
    # --------------------------------------------------------

    log_records = []


    # --------------------------------------------------------
    # Process every student
    # --------------------------------------------------------

    for index, row in df.head(1).iterrows():

        year = str(row["Year"]).strip()

        name = str(row["Name"]).strip()
        academic_year = f"{year}-{str(int(year) + 1)[-2:]}"
        application_number = (
            str(row["Application Number"]).strip()
        )

        rank = str(row["Rank"]).strip()

        score = str(row["Score"]).strip()


        print(
            f"\n[{index + 1}/{len(df)}] "
            f"{name}"
        )


        # ----------------------------------------------------
        # Validate record
        # ----------------------------------------------------

        if not name or name.lower() == "nan":

            print("   ❌ Invalid name")

            log_records.append({

                "Year": year,

                "Name": name,

                "Application Number":
                    application_number,

                "Rank": rank,

                "Score": score,

                "Status": "FAILED",

                "PDF File": "",

                "Message": "Invalid name"
            })

            continue


        # ----------------------------------------------------
        # Check if PDF already exists
        # ----------------------------------------------------

        safe_name = (
            name
            .replace("/", "_")
            .replace("\\", "_")
            .replace(" ", "_")
        )


        normal_pdf = (
            OUTPUT_DIR
            / f"{safe_name}_Scorecard.pdf"
        )


        duplicate_pdf = (
            OUTPUT_DIR
            / f"{safe_name}_{application_number}_Scorecard.pdf"
        )


        if normal_pdf.exists() or duplicate_pdf.exists():

            existing_file = (
                normal_pdf
                if normal_pdf.exists()
                else duplicate_pdf
            )


            print(
                "   ⏭️ PDF already exists - skipped"
            )


            log_records.append({

                "Year": year,

                "Name": name,

                "Application Number":
                    application_number,

                "Rank": rank,

                "Score": score,

                "Status": "SKIPPED",

                "PDF File":
                    existing_file.name,

                "Message":
                    "PDF already exists"
            })

            continue


        # ----------------------------------------------------
        # Generate PDF
        # ----------------------------------------------------

        try:

            pdf_file = create_scorecard(

                year=year,
                academic_year=academic_year,
                name=name,

                application_number=
                    application_number,

                rank=rank,

                score=score
            )


            print(
                f"   ✅ Generated: "
                f"{pdf_file.name}"
            )


            log_records.append({

                "Year": year,

                "Name": name,

                "Application Number":
                    application_number,

                "Rank": rank,

                "Score": score,

                "Status": "GENERATED",

                "PDF File":
                    pdf_file.name,

                "Message": "Success"
            })


        except Exception as error:

            print(
                f"   ❌ Failed: {error}"
            )


            log_records.append({

                "Year": year,

                "Name": name,

                "Application Number":
                    application_number,

                "Rank": rank,

                "Score": score,

                "Status": "FAILED",

                "PDF File": "",

                "Message": str(error)
            })


    # --------------------------------------------------------
    # Save generation log
    # --------------------------------------------------------

    log_df = pd.DataFrame(
        log_records
    )


    log_df.to_excel(
        LOG_FILE,
        index=False
    )


    # --------------------------------------------------------
    # Final summary
    # --------------------------------------------------------

    generated_count = len(
        log_df[
            log_df["Status"] == "GENERATED"
        ]
    )


    skipped_count = len(
        log_df[
            log_df["Status"] == "SKIPPED"
        ]
    )


    failed_count = len(
        log_df[
            log_df["Status"] == "FAILED"
        ]
    )


    print("\n" + "=" * 70)
    print("BULK GENERATION COMPLETE")
    print("=" * 70)


    print(
        f"Total Records : {len(df)}"
    )

    print(
        f"Generated     : {generated_count}"
    )

    print(
        f"Skipped       : {skipped_count}"
    )

    print(
        f"Failed        : {failed_count}"
    )


    print(
        f"\nPDF Folder:"
        f"\n{OUTPUT_DIR}"
    )


    print(
        f"\nGeneration Log:"
        f"\n{LOG_FILE}"
    )


# ============================================================
# RUN
# ============================================================

if __name__ == "__main__":

    generate_all_scorecards()